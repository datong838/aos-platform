#!/usr/bin/env python3
"""
通用专题 Word 文档生成器
用法: python generate_topic.py <专题目录路径> [--output-dir 输出目录]

示例:
  python generate_topic.py /Users/ddt/work/projects/ai_agent/docs/palantier/foundry/pages/zh/foundry/action-types
  python generate_topic.py /path/to/topic --output-dir /path/to/output
"""
import os
import re
import sys
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

# ============================================================
# 默认配置
# ============================================================
DEFAULT_OUTPUT_DIR = Path("/Users/ddt/work/projects/ai_agent/docs/palantier/prddetail")
PALANTIER_ROOT = Path("/Users/ddt/work/projects/ai_agent/docs/palantier")


def scan_all_images(topic_name):
    """
    扫描所有可能的图片目录，建立文件名→路径的全局索引。
    覆盖 images/ 和 resources/ 两个目录树。
    """
    all_images = {}
    search_dirs = [
        (PALANTIER_ROOT / "foundry" / "images" / "foundry" / topic_name),
        (PALANTIER_ROOT / "resources" / "foundry" / topic_name),
    ]
    for search_dir in search_dirs:
        if search_dir.exists():
            for img_path in search_dir.rglob("*"):
                if img_path.is_file() and img_path.suffix.lower() in (
                    '.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'
                ):
                    all_images[img_path.name] = img_path
    return all_images


def discover_doc_order(pages_dir):
    """
    发现文档顺序。优先使用 frontmatter 中的 previous/next 链，
    其次按字母序排列，但 overview 始终排第一。
    """
    md_files = sorted([
        f for f in pages_dir.iterdir()
        if f.suffix == '.md' and f.is_file()
    ], key=lambda x: x.name)

    if not md_files:
        return []

    # 尝试构建 previous/next 链
    links = {}  # name -> {prev, next}
    for f in md_files:
        with open(f, 'r', encoding='utf-8') as fh:
            content = fh.read()
        fm = {}
        if content.startswith("---"):
            idx = content.find("---", 3)
            if idx != -1:
                try:
                    import json
                    fm = json.loads(content[3:idx].strip())
                except (json.JSONDecodeError, ValueError):
                    # 尝试用简单的 key: value 解析
                    for line in content[3:idx].strip().split('\n'):
                        line = line.strip()
                        if ':' in line:
                            k, v = line.split(':', 1)
                            fm[k.strip().strip('"')] = v.strip().strip('"')
        links[f.name] = {
            'prev': fm.get('previous', ''),
            'next': fm.get('next', ''),
            'title': fm.get('title', ''),
        }

    # 找起点（没有 prev 的，或者 overview）
    ordered = []
    visited = set()
    remaining = set(f.name for f in md_files)

    # overview 优先
    overview_candidates = [n for n in remaining if 'overview' in n.lower()]
    for ov in overview_candidates:
        remaining.discard(ov)

    # 构建链接顺序
    def follow_chain(start_name):
        chain = []
        current = start_name
        while current and current in remaining:
            chain.append(current)
            visited.add(current)
            remaining.discard(current)
            nxt = links.get(current, {}).get('next', '')
            if nxt in remaining and nxt not in visited:
                current = nxt
            else:
                break
        return chain

    # 从每个可能的起始点 follow
    starters = [n for n in remaining if not links.get(n, {}).get('prev', '')]
    for s in starters:
        ordered.extend(follow_chain(s))

    # 剩余未连接的按字母序
    leftover = sorted(remaining)
    ordered.extend(leftover)

    # overview 放最前
    ordered = overview_candidates + [f for f in ordered if f not in overview_candidates]

    # 确保所有文件都在列表中
    for f in md_files:
        if f.name not in ordered:
            ordered.append(f.name)

    return ordered


def resolve_image_path(img_ref, current_file_dir, all_images, images_dir):
    """
    多策略解析图片引用到实际文件路径。
    策略优先级：直接文件检查 → 相对路径 → /resources/ 绝对路径 → 全局文件名匹配
    """
    img_ref = img_ref.strip()

    if not img_ref:
        return None

    # 策略 1: 如果是 / 开头的绝对路径（如 /resources/...），从 PALANTIER_ROOT 拼接
    if img_ref.startswith('/'):
        abs_path = PALANTIER_ROOT / img_ref.lstrip('/')
        if abs_path.exists():
            return str(abs_path)

    # 策略 2: 相对路径解析
    resolved = (Path(current_file_dir) / img_ref).resolve()
    if resolved.exists():
        return str(resolved)

    # 策略 3: 从图片目录按文件名查找（覆盖了跨目录的相对路径情况）
    filename = Path(img_ref).name
    if filename in all_images:
        return str(all_images[filename])

    # 策略 4: 从 images_dir 直接查找
    try_path = images_dir / filename
    if try_path.exists():
        return str(try_path)

    return None


def parse_frontmatter(content):
    """解析 YAML/JSON frontmatter"""
    if not content.startswith("---"):
        return {}, content
    idx = content.find("---", 3)
    if idx == -1:
        return {}, content
    try:
        import json
        fm = json.loads(content[3:idx].strip())
    except (json.JSONDecodeError, ValueError):
        fm = {}
    return fm, content[idx + 3:]


def strip_callout_blocks(content):
    """移除 :::callout 块标记，保留内容文本"""
    pattern = r':::callout\{[^}]*\}(?:.*?\n)?([\s\S]*?):::\n?'
    content = re.sub(pattern, lambda m: m.group(1).strip() + "\n\n", content)
    content = re.sub(r':::callout\{[^}]*\}', '', content)
    content = re.sub(r'^:::\s*$', '', content, flags=re.MULTILINE)
    return content


def insert_image_from_markdown(doc, img_path, alt_text="", max_width=5.5):
    """插入图片，自动计算合适尺寸"""
    if not img_path or not Path(img_path).exists():
        p = doc.add_paragraph()
        run = p.add_run(f'[图片缺失：{alt_text or Path(img_path).name if img_path else "unknown"}]')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 100, 100)
        run.italic = True
        return False

    try:
        img = Image.open(img_path)
        img_width, img_height = img.size
        aspect = img_height / img_width if img_width > 0 else 1
        width_inches = min(max_width, 6.0)
        height_inches = width_inches * aspect

        max_height = 4.5
        if height_inches > max_height:
            height_inches = max_height
            width_inches = height_inches / aspect if aspect > 0 else max_height

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Inches(width_inches), height=Inches(height_inches))
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(12)
        return True
    except Exception as e:
        p = doc.add_paragraph()
        run = p.add_run(f'[图片加载失败：{alt_text or Path(img_path).name}]')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 0, 0)
        return False


def process_inline_text(paragraph, text):
    """处理内联格式化：粗体、代码、链接"""
    pattern = r'(\*\*(.+?)\*\*|`(.+?)`|\[([^\]]+)\]\(([^)]+)\)|\\\[|\]\()'
    last_end = 0
    for match in re.finditer(pattern, text):
        if match.start() > last_end:
            run = paragraph.add_run(text[last_end:match.start()])
            run.font.name = '微软雅黑'
            run.font.size = Pt(10.5)

        full = match.group(0)
        if match.group(2):  # bold
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.font.name = '微软雅黑'
            run.font.size = Pt(10.5)
        elif match.group(3):  # code
            run = paragraph.add_run(match.group(3))
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(180, 60, 60)
        elif match.group(4):  # link
            run = paragraph.add_run(match.group(4))
            run.font.name = '微软雅黑'
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0, 102, 204)
            run.underline = True
        elif full in ['\\[', '](']:
            run = paragraph.add_run(full)
            run.font.name = '微软雅黑'
            run.font.size = Pt(10.5)

        last_end = match.end()

    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        run.font.name = '微软雅黑'
        run.font.size = Pt(10.5)

    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(6)


def process_single_page(doc, filepath, all_images, images_dir):
    """处理单个 markdown 文件，写入 Word"""
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()

    fm, body = parse_frontmatter(content)
    title = fm.get("title", "")
    body = strip_callout_blocks(body)
    current_file_dir = filepath.parent

    if title:
        h = doc.add_heading(title, level=1)
        for run in h.runs:
            run.font.name = '微软雅黑'

    lines = body.split('\n')
    i = 0
    table_rows = []
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        # 空行
        if not line.strip():
            i += 1
            continue

        # 代码块
        if line.strip().startswith('```'):
            if in_code_block:
                code_text = '\n'.join(code_lines)
                if code_text.strip():
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(1)
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(4)
                    run = p.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(80, 80, 80)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # 标题
        if line.startswith('## '):
            h = doc.add_heading(line[3:].strip(), level=2)
            for run in h.runs:
                run.font.name = '微软雅黑'
            i += 1
            continue

        if line.startswith('### '):
            h = doc.add_heading(line[4:].strip(), level=3)
            for run in h.runs:
                run.font.name = '微软雅黑'
            i += 1
            continue

        # 水平线
        if line.strip() in ('***', '---'):
            doc.add_paragraph().add_run('─' * 50)
            i += 1
            continue

        # === 图片：Markdown 格式 ===
        img_md = re.search(r'!\[([^\]]*)\]\(([^)]+)\)', line)
        if img_md:
            alt_text = img_md.group(1)
            img_ref = img_md.group(2)

            prefix = line[:img_md.start()].strip()
            if prefix:
                list_match = re.match(r'^(\s*)[\*\-\+]\s+(.*)', prefix)
                num_match = re.match(r'^(\s*)\d+\.\s+(.*)', prefix)
                if list_match:
                    p = doc.add_paragraph(style='List Bullet')
                    process_inline_text(p, list_match.group(2))
                elif num_match:
                    p = doc.add_paragraph(style='List Number')
                    process_inline_text(p, num_match.group(2))
                else:
                    p = doc.add_paragraph()
                    process_inline_text(p, prefix)

            # 图片标题
            if alt_text:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f'图：{alt_text}')
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(100, 100, 100)
                run.italic = True
                p.paragraph_format.space_after = Pt(2)

            resolved_path = resolve_image_path(img_ref, current_file_dir, all_images, images_dir)
            if resolved_path:
                insert_image_from_markdown(doc, resolved_path, alt_text)
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f'[图片未找到：{img_ref}]')
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 100, 100)
                run.italic = True

            suffix = line[img_md.end():].strip()
            if suffix:
                p = doc.add_paragraph()
                process_inline_text(p, suffix)

            i += 1
            continue

        # === 图片：HTML 格式 ===
        img_html = re.search(r'<img\s+src="([^"]+)"[^>]*>', line)
        if img_html:
            img_ref = img_html.group(1)
            resolved_path = resolve_image_path(img_ref, current_file_dir, all_images, images_dir)
            if resolved_path:
                insert_image_from_markdown(doc, resolved_path)
            i += 1
            continue

        # === 表格 ===
        if line.strip().startswith('|') and line.strip().endswith('|'):
            if re.match(r'^\|[\s\-:|]+\|$', line.strip()):
                i += 1
                continue

            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            table_rows.append(cells)

            if (i + 1 >= len(lines) or
                not (lines[i + 1].strip().startswith('|') and lines[i + 1].strip().endswith('|'))):
                if table_rows:
                    cols = len(table_rows[0])
                    table = doc.add_table(rows=len(table_rows), cols=cols, style='Light Grid Accent 1')
                    for ri, row_data in enumerate(table_rows):
                        for ci, cell_data in enumerate(row_data[:cols]):
                            cell = table.rows[ri].cells[ci]
                            cell.text = cell_data
                            for p_in_cell in cell.paragraphs:
                                for run in p_in_cell.runs:
                                    run.font.size = Pt(10)
                                    run.font.name = '微软雅黑'
                                p_in_cell.paragraph_format.space_after = Pt(2)
                                p_in_cell.paragraph_format.space_before = Pt(2)
                    doc.add_paragraph()
                table_rows = []
            i += 1
            continue

        # === 无序列表 ===
        ul_match = re.match(r'^(\s*)[\*\-\+]\s+(.+)', line)
        if ul_match:
            p = doc.add_paragraph(style='List Bullet')
            process_inline_text(p, ul_match.group(2))
            i += 1
            continue

        # === 有序列表 ===
        ol_match = re.match(r'^(\s*)\d+\.\s+(.+)', line)
        if ol_match:
            p = doc.add_paragraph(style='List Number')
            process_inline_text(p, ol_match.group(2))
            i += 1
            continue

        # === 引用块 ===
        if line.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run('▎ ')
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(150, 150, 150)
            run2 = p.add_run(line[2:])
            run2.font.size = Pt(10)
            run2.font.color.rgb = RGBColor(100, 100, 100)
            run2.italic = True
            i += 1
            continue

        # === 粗体标记行（如 "**标题**：内容"） ===
        bold_match = re.match(r'^\*\*(.+?)\*\*\s*[：:]?\s*(.*)', line)
        if bold_match:
            p = doc.add_paragraph()
            run = p.add_run(bold_match.group(1))
            run.bold = True
            run.font.name = '微软雅黑'
            if bold_match.group(2):
                run2 = p.add_run(f'：{bold_match.group(2)}')
                run2.font.name = '微软雅黑'
            i += 1
            continue

        # === 普通段落 ===
        p = doc.add_paragraph()
        process_inline_text(p, line)
        i += 1


def get_topic_display_name(topic_name):
    """将目录名转换成可读的显示名称"""
    # action-types → Action Types
    return topic_name.replace('-', ' ').replace('_', ' ').title()


def generate_topic_doc(topic_dir, output_dir, topic_name=None):
    """主生成函数"""
    topic_dir = Path(topic_dir)
    output_dir = Path(output_dir)

    if not topic_dir.exists():
        print(f"❌ 专题目录不存在: {topic_dir}")
        sys.exit(1)

    # 自动推断专题名称（目录的最后一级）
    if topic_name is None:
        topic_name = topic_dir.name
    display_name = get_topic_display_name(topic_name)

    # 图片目录
    images_dir = PALANTIER_ROOT / "foundry" / "images" / "foundry" / topic_name
    all_images = scan_all_images(topic_name)

    print(f"📂 专题目录: {topic_dir}")
    print(f"📂 图片目录: {images_dir}")
    print(f"🖼  图片索引: {len(all_images)} 张")

    # 发现文档列表
    doc_order = discover_doc_order(topic_dir)
    print(f"📄 发现文档: {len(doc_order)} 个")

    output_dir.mkdir(parents=True, exist_ok=True)

    # ========================================
    # 构建 Word 文档
    # ========================================
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.5)

    # Normal 样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.5

    # Heading 1
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = '微软雅黑'
    h1_style.font.size = Pt(18)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 51, 102)
    h1_style.paragraph_format.space_before = Pt(24)
    h1_style.paragraph_format.space_after = Pt(12)

    # Heading 2
    h2_style = doc.styles['Heading 2']
    h2_style.font.name = '微软雅黑'
    h2_style.font.size = Pt(15)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0, 70, 127)
    h2_style.paragraph_format.space_before = Pt(18)
    h2_style.paragraph_format.space_after = Pt(8)

    # Heading 3
    h3_style = doc.styles['Heading 3']
    h3_style.font.name = '微软雅黑'
    h3_style.font.size = Pt(13)
    h3_style.font.bold = True
    h3_style.font.color.rgb = RGBColor(0, 90, 150)
    h3_style.paragraph_format.space_before = Pt(12)
    h3_style.paragraph_format.space_after = Pt(6)

    # ========================================
    # 封面
    # ========================================
    for _ in range(6):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(topic_name)
    run.font.size = Pt(32)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    run.font.name = '微软雅黑'

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run(display_name)
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(80, 80, 80)
    run.font.name = '微软雅黑'

    doc.add_paragraph()
    doc.add_paragraph()

    desc_p = doc.add_paragraph()
    desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc_p.add_run(f'Palantir Foundry {display_name} 完整参考文档')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.font.name = '微软雅黑'

    doc.add_paragraph()

    source_p = doc.add_paragraph()
    source_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = source_p.add_run('基于 Palantir Foundry 官方文档整理')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.font.name = '微软雅黑'

    doc.add_page_break()

    # ========================================
    # 目录
    # ========================================
    toc_title = doc.add_heading('目录', level=1)

    for fname in doc_order:
        filepath = topic_dir / fname
        if not filepath.exists():
            continue
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        fm, _ = parse_frontmatter(content)
        fm_title = fm.get("title", fname.replace('.md', '').replace('-', ' ').title())

        p = doc.add_paragraph()
        run = p.add_run(fm_title)
        run.font.name = '微软雅黑'
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ========================================
    # 正文
    # ========================================
    found_count = 0
    missing_count = 0

    for idx, fname in enumerate(doc_order):
        filepath = topic_dir / fname
        if not filepath.exists():
            print(f"  ⚠ 文件不存在: {fname}")
            continue

        print(f"  📄 {fname}")

        if idx > 0:
            doc.add_page_break()

        # 统计图片
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        for img_match in re.finditer(r'!\[([^\]]*)\]\(([^)]+)\)', content):
            img_ref = img_match.group(2)
            resolved = resolve_image_path(img_ref, filepath.parent, all_images, images_dir)
            if resolved:
                found_count += 1
            else:
                missing_count += 1
                print(f"    ⚠ 图片未找到: {img_ref}")

        process_single_page(doc, filepath, all_images, images_dir)

    # ========================================
    # 保存
    # ========================================
    output_path = output_dir / f"{topic_name}.docx"
    doc.save(str(output_path))

    print(f"\n✅ 文档已生成: {output_path}")
    print(f"  📊 处理文档: {len(doc_order)} 个")
    print(f"  🖼  嵌入图片: {found_count} 张")
    if missing_count:
        print(f"  ⚠ 缺失图片: {missing_count} 张")

    return output_path


def main():
    parser = argparse.ArgumentParser(
        description='从 Markdown 专题目录生成标准 PRD 格式的 Word 文档',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  %(prog)s /path/to/topic/pages
  %(prog)s /path/to/topic/pages --output-dir /path/to/output
  %(prog)s /path/to/topic/pages --topic-name my-topic
        """
    )
    parser.add_argument('topic_dir', help='专题 Markdown 文件所在目录')
    parser.add_argument('--output-dir', '-o', type=str, default=str(DEFAULT_OUTPUT_DIR),
                        help=f'输出目录（默认: {DEFAULT_OUTPUT_DIR}）')
    parser.add_argument('--topic-name', '-n', type=str, default=None,
                        help='专题名称，用于文件名和封面（默认取目录名）')

    args = parser.parse_args()

    generate_topic_doc(
        topic_dir=args.topic_dir,
        output_dir=args.output_dir,
        topic_name=args.topic_name,
    )


if __name__ == '__main__':
    main()
